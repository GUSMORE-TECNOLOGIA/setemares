from __future__ import annotations

from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx2pdf import convert


def render_from_docx(template_path: str, out_docx: str, out_pdf: str, data: Dict[str, Any]) -> None:
	tpl = Document(template_path)
	# Substituição simples de placeholders no corpo
	from datetime import datetime
	now = datetime.now()
	data = {**data, "data_emissao": now.strftime("%d/%m/%Y"), "hora_emissao": now.strftime("%H:%M")}
	for p in tpl.paragraphs:
		for k, v in data.items():
			placeholder = f"{{{{{k}}}}}"
			if placeholder in p.text:
				p.text = p.text.replace(placeholder, str(v))
	# Tabelas (primeira tabela como grade de voos, se existir)
	if tpl.tables and data.get("decoded"):
		flights = data["decoded"].get("flightInfo", {}).get("flights", [])
		tab = tpl.tables[0]
		# Limpa linhas além do cabeçalho (mantém primeira)
		while len(tab.rows) > 1:
			_tab_row = tab.rows[-1]
			tab._tbl.remove(_tab_row._tr)
		for f in flights:
			row = tab.add_row()
			row.cells[0].text = f"{f['company']['iataCode']} {f['flight']}"
			row.cells[1].text = f["departureAirport"].get("description") or f["departureAirport"].get("iataCode")
			row.cells[2].text = f["landingAirport"].get("description") or f["landingAirport"].get("iataCode")
			row.cells[3].text = f.get("departureTime", "")
			row.cells[4].text = f.get("landingTime", "")
		tab.style = tab.style or 'Table Grid'

	# Garante diretório do DOCX temporário
	from pathlib import Path as _P
	_P(out_docx).parent.mkdir(parents=True, exist_ok=True)
	tpl.save(out_docx)
	# Converter para PDF usando Word (docx2pdf)
	convert(out_docx, out_pdf)
